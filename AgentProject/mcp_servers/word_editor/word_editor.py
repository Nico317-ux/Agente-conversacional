import pythoncom
from win32com.client import Dispatch
from typing import Optional, Dict
from mcp.server.fastmcp import FastMCP
from pathlib import Path
from docx.shared import Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from utils_word.word_utils import format_json, DocxEditor
from docx.oxml.ns import qn

mcp = FastMCP('docx_processor')
doc_editor = DocxEditor()

@mcp.tool()
async def create_docx(margins: Optional[Dict[str, float]] = None) -> str:
    '''Create a new Word document and return its path.
          
        Args:
            margins (dict): Optional dictionary with margin settings in Cm.
                Example: {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0}

        Retruns:
              JSON response with:
              -'status': ''success' or 'error'
              -'message': 'Document created successfully' or error message
    '''
    try:
        default_margins = {
            'top': 1.0,
            'bottom': 1.0,
            'left': 1.0,
            'right': 1.0
        }
        final_margins = margins or default_margins
        doc_editor.create_new_docx(final_margins)
        return format_json(status='success', message='Document created successfully')
    except Exception as e:
        return format_json(status='error', message=str(e))
    
@mcp.tool()
async def open_docx(file_path: str) -> str:
    '''Open an existing Word document and return its path.
          
        Args:
            file_path (str): The path to the Word document to open.
          
        Returns:
            JSON response with:
            -'status': 'success' or 'error'
            -'message': 'Document opened successfully' or error message
            -'path': The path to the opened document
    '''
    try:
        path = Path(file_path)
        if not path.exists():
            return format_json(status='error', message='File does not exist')
        doc_editor.open_docx(file_path)
        return format_json(status='success', message='Document opened successfully', path=file_path)
    except Exception as e:
        return format_json(status='error', message=str(e))
    
@mcp.tool()
async def save_docx(file_path: Optional[str] = None) -> str:
    '''Save the current Word document to the specified path.
          
        Args:
            file_path (str): The path where the document should be saved.
          
        Returns:
            JSON response with:
            -'status': 'success' or 'error'
            -'message': 'Document saved successfully' or error message
    '''
    try:
        if not doc_editor.current_docx:
            return format_json(status='error', message='No document is currently open')
        
        save_path = file_path or doc_editor.doc_path
        if not save_path:
            return format_json(status='error', message='No save path specified')
        doc_editor.current_docx.save(save_path)
        doc_editor.doc_path = save_path
        return format_json(status='success', message='Document saved successfully', path=str(save_path))
    except Exception as e:
        return format_json(status='error', message=str(e))
    
@mcp.tool()
async def add_paragraph(
    text: str,
    font_size: Optional[int] = 12,
    font_name: Optional[str] = 'Times New Roman',
    bold: Optional[bool] = False,
    italic: Optional[bool] = False,
    aligment: Optional[str] = 'justify', # left/center/right/justify
    spacing_before: Optional[float] = 0,
    spacing_after: Optional[float] = 12,
    line_spacing: Optional[float] = 1.5,
    first_line_indent: Optional[float] = None
) -> str:
    '''Add a paragraph to the current Word document with specified formatting.
          
        Args:
            text (str): The text of the paragraph to add.
            font_size (int): The font size of the paragraph text.
            bold (bool): Whether the text should be bold.
            italic (bool): Whether the text should be italicized.
            aligment (str): The alignment of the paragraph ('left', 'center', 'right', 'justify').
            spacing_before (float): Space before the paragraph in points.
            spacing_after (float): Space after the paragraph in points.
            line_spacing (float): Line spacing for the paragraph.
          
        Returns:
            JSON response with:
            -'status': 'success' or 'error'
            -'message': 'Paragraph added successfully' or error message
            -'data': {"format": {
                "font_size": font_size,
                "bold": bold,
                "italic": italic,
                "alignment": alignment,
                "spacing_before": spacing_before,
                "spacing_after": spacing_after,
                "line_spacing": line_spacing
            }}
    '''
    try:
        if not doc_editor.current_docx:
            return format_json(status='error', message='No document is currently open')
        paragraph = doc_editor.current_docx.add_paragraph(text)
        run = paragraph.runs[0]

        font = run.font
        font.size = Pt(font_size)
        font.name = font_name

        font._element.rPr.rFonts.set(qn('w:TimesNewRoman'), font_name)
        font.bold = bold
        font.italic = italic

        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        paragraph.alignment = align_map.get(aligment, WD_PARAGRAPH_ALIGNMENT.LEFT)

        p_format = paragraph.paragraph_format
        p_format.space_before = Pt(spacing_before)
        p_format.space_after = Pt(spacing_after)
        p_format.line_spacing = line_spacing

        if first_line_indent:
            p_format.first_line_indent = Pt(first_line_indent)

        return format_json(
            status='success',
            message='Paragraph added successfully',
            data={
                "format": {
                    "font_name": font_name,
                    "font_size": font_size,
                    "bold": bold,
                    "italic": italic,
                    "alignment": aligment,
                    "spacing_before": spacing_before,
                    "spacing_after": spacing_after,
                    "line_spacing": line_spacing,
                    "first_line_indent": first_line_indent
                }
            }
        )
    except Exception as e:
        return format_json(status='error', message=str(e))
    
@mcp.tool()
async def add_heading(text: str, level: int = 1) -> str:
    '''Add a heading to the current Word document.
          
        Args:
            text (str): The text of the heading to add.
            level (int): The level of the heading (1-9).
          
        Returns:
            JSON response with:
            -'status': 'success' or 'error'
            -'message': 'Heading added successfully' or error message
    '''
    try:
        if not doc_editor.current_docx:
            raise ValueError("No document is currently open")
        doc_editor.current_docx.add_heading(text, level=level)
        return format_json(status='success', message='Heading added successfully')
    except Exception as e:
        return format_json(status='error', message=str(e))

@mcp.tool()
async def add_image(image_path: str, width: float = 6.0) -> str:
    '''Add an image to the current Word document.

        Args:
            image_path (str): The path to the image file.
            width (float): The width of the image in Cm (default is 6.0).
        
        Returns:
            JSON response with:
            -'status': 'success' or 'error'
            -'message': 'Image added successfully' or error message
            -'path': The path to the added image
    '''
    try:
        path = Path(image_path)
        if not path.exists():
            raise ValueError("the image file does not exist")
        doc_editor.current_docx.add_picture(str(path), width=Cm(width))
        return format_json(status='success', message='Image added successfully', path=str(path))
    except Exception as e:
        return format_json(status='error', message=str(e))
 
@mcp.tool()
async def add_table(rows: int, cols: int, data: Optional[list] = None, style: Optional[str] = 'Table Grid') -> str:
    '''Add a table to the current Word document.
          
        Args:
            rows (int): Number of rows in the table.
            cols (int): Number of columns in the table.
            data (list): Optional 2D list of data to fill the table.
          
        Returns:
            JSON response with:
            -'status': 'success' or 'error'
            -'message': 'Table added successfully' or error message
    '''
    try:
        table = doc_editor.current_docx.add_table(rows = rows, cols= cols)
        table.style = style or 'Table Grid'

        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                row = table.rows[i]
                for j, cell_data in enumerate(row_data):
                    if j >= cols:
                        break
                    row.cells[j].text = str(cell_data)
        return format_json(status='success', message='Table added successfully')
    except Exception as e:
        return format_json(status='error', message=str(e))

@mcp.tool()
async def preview_docx() -> str:
    '''Open the current Word document in Microsoft Word for preview.
          
        Returns:
            JSON response with:
            -'status': 'success' or 'error'
            -'message': 'Document opened for preview' or error message
    '''
    try:
        if not doc_editor.doc_path:
            raise ValueError("You must save the document first")
        pythoncom.CoInitialize()
        word = Dispatch("Word.Application")
        word.Visible = True
        word.Documents.Open(str(doc_editor.doc_path))

        return format_json(status='success', message='Document opened for preview', path=str(doc_editor.doc_path))
    except Exception as e:
        return format_json(status='error', message=str(e))

if __name__ == '__main__':
    mcp.run(transport='stdio')